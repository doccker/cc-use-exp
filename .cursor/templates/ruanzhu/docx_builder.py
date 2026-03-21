"""
DOCX 文档构建模块

职责：渲染行数估算、源码行构建、DOCX 文件生成
"""

from code_scanner import read_file_lines

# DOCX格式配置
DOCX_CONFIG = {
    'page_width_cm': 21,
    'page_height_cm': 29.7,
    'margin_top_cm': 2.5,
    'margin_bottom_cm': 2.5,
    'margin_left_cm': 3.0,
    'margin_right_cm': 2.5,
    'font_size_pt': 10,
    'font_cn': '宋体',
    'font_en': 'Courier New',
    'lines_per_page': 54,
}


def estimate_rendered_lines(text, usable_width_cm=15.5):
    """估算一行文本在Word中渲染需要的行数"""
    if len(text) == 0:
        return 1
    width = 0
    for ch in text:
        if ord(ch) > 127:
            width += 0.35
        else:
            width += 0.18
    return max(1, int(width / usable_width_cm) + (1 if width % usable_width_cm > 0.1 else 0))


def build_source_lines(source_files, target_rendered_lines, mode='fixed'):
    """
    构建源代码行

    Args:
        source_files: 源文件列表 [(filename, lang, path), ...]
        target_rendered_lines: 目标渲染行数
        mode: 'fixed' 固定页数, 'auto' 自动模式
    """
    all_lines = []
    total_rendered = 0

    for filename, _lang, filepath in source_files:
        lines = read_file_lines(filepath)
        if not lines:
            continue

        header_line = f"/* ========== {filename} ========== */"
        header_rendered = estimate_rendered_lines(header_line)

        if mode == 'fixed' and total_rendered + header_rendered + 2 > target_rendered_lines:
            break

        all_lines.append(header_line)
        all_lines.append("")
        total_rendered += header_rendered + 1

        for line in lines:
            line_rendered = estimate_rendered_lines(line)
            if mode == 'fixed' and total_rendered + line_rendered > target_rendered_lines:
                break
            all_lines.append(line)
            total_rendered += line_rendered

        all_lines.append("")
        total_rendered += 1

        if mode == 'fixed' and total_rendered >= target_rendered_lines:
            break

    return all_lines, total_rendered


def _add_field_with_text(para, field_code, placeholder, cfg):
    """添加 Word 域代码，包含占位文本以确保正确渲染"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f" {field_code} "
    run2._element.append(instrText)

    run3 = para.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar2)

    run4 = para.add_run(placeholder)
    run4.font.name = cfg['font_en']
    run4.font.size = Pt(cfg['font_size_pt'])
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

    run5 = para.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar3)


def _add_styled_run(para, text, cfg):
    """添加带统一样式的文本 run"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run = para.add_run(text)
    run.font.name = cfg['font_en']
    run.font.size = Pt(cfg['font_size_pt'])
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])
    return run


def _add_code_paragraph(doc, line, cfg):
    """添加一行代码段落"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(line)
    run.font.size = Pt(cfg['font_size_pt'])
    run.font.name = cfg['font_en']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])


def _setup_header(section, software_name, version, cfg):
    """设置页眉：左侧标题 + 居中页码"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0]

    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_center = OxmlElement('w:tab')
    tab_center.set(qn('w:val'), 'center')
    tab_center.set(qn('w:pos'), '4394')
    tabs.append(tab_center)
    pPr.append(tabs)

    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    run1 = _add_styled_run(p, f"{software_name} {version}", cfg)
    run1.font.bold = True

    run_tab = p.add_run()
    tab_elem = OxmlElement('w:tab')
    run_tab._element.append(tab_elem)

    _add_styled_run(p, "第 ", cfg)
    _add_field_with_text(p, "PAGE", "1", cfg)
    _add_styled_run(p, " 页共 ", cfg)
    _add_field_with_text(p, "NUMPAGES", "1", cfg)
    _add_styled_run(p, " 页", cfg)


def _setup_footer(section, cfg):
    """设置页脚：居中页码"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, "第 ", cfg)
    _add_field_with_text(p, "PAGE", "1", cfg)
    _add_styled_run(p, " 页", cfg)


def generate_docx(lines, output_path, software_name, version, is_split=False, total_pages=0):
    """生成DOCX文件"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    cfg = DOCX_CONFIG

    section = doc.sections[0]
    section.page_width = Cm(cfg['page_width_cm'])
    section.page_height = Cm(cfg['page_height_cm'])
    section.top_margin = Cm(cfg['margin_top_cm'])
    section.bottom_margin = Cm(cfg['margin_bottom_cm'])
    section.left_margin = Cm(cfg['margin_left_cm'])
    section.right_margin = Cm(cfg['margin_right_cm'])

    _setup_header(section, software_name, version, cfg)
    _setup_footer(section, cfg)

    style = doc.styles['Normal']
    style.font.size = Pt(cfg['font_size_pt'])
    style.font.name = cfg['font_en']
    style.element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    if is_split:
        front_lines, back_lines = lines
        for line in front_lines:
            _add_code_paragraph(doc, line, cfg)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"\n（以上为源代码前30页，以下为源代码后30页，中间省略部分共{total_pages - 60}页）\n"
        )
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = cfg['font_en']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

        for line in back_lines:
            _add_code_paragraph(doc, line, cfg)
    else:
        for line in lines:
            _add_code_paragraph(doc, line, cfg)

    doc.save(output_path)
    return output_path
